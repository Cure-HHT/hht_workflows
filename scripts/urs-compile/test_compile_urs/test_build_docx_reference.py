"""Smoke tests for build_docx_reference.py — exercise the style-builder so
regressions in heading sizes, table styling, or custom-style definitions
are caught without a manual Word/LibreOffice inspection."""
from __future__ import annotations

import sys
import zipfile
from pathlib import Path

import pytest
from lxml import etree

SCRIPT_DIR = Path(__file__).parents[1]
sys.path.insert(0, str(SCRIPT_DIR))

import build_docx_reference  # noqa: E402

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


@pytest.fixture
def built_docx(tmp_path):
    out = tmp_path / "ref.docx"
    build_docx_reference.build_reference_docx(
        out,
        {
            "sponsor_name": "Test Sponsor",
            "protocol_number": "TST-0001",
            "protocol_version": "1.0",
        },
    )
    return out


def _styles_root(path: Path):
    with zipfile.ZipFile(path) as z:
        return etree.fromstring(z.read("word/styles.xml"))


def _find_style(root, display_name: str):
    for st in root.findall(f"{{{W}}}style"):
        nm = st.find(f"{{{W}}}name")
        if nm is not None and nm.get(f"{{{W}}}val") == display_name:
            return st
    return None


def test_heading_1_through_3_keep_urs_blue_and_sizes(built_docx):
    from docx import Document
    d = Document(built_docx)
    by_name = {s.name: s for s in d.styles}
    for name, expected_size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        s = by_name[name]
        assert s.font.size.pt == expected_size
        assert s.font.bold is True
        assert str(s.font.color.rgb) == "1F3A5F"
    # Heading 1 (chapter) and Heading 3 (REQ) both start fresh pages so the
    # docx mirrors the LaTeX template's \newpage on \chapter and \subsection.
    assert by_name["Heading 1"].paragraph_format.page_break_before is True
    assert by_name["Heading 3"].paragraph_format.page_break_before is True
    assert by_name["Heading 2"].paragraph_format.page_break_before is not True


def test_heading_4_and_5_are_subheading_styled(built_docx):
    """H4 / H5 are bold + URS-blue (smaller than chapter/section headings)
    so REQ-internal subheadings stand out, but have outlineLvl cleared so
    they don't appear in Word's document outline or the TOC."""
    from docx import Document
    d = Document(built_docx)
    by_name = {s.name: s for s in d.styles}
    for name, expected_size in (("Heading 4", 11), ("Heading 5", 10)):
        s = by_name[name]
        assert s.font.size.pt == expected_size
        assert s.font.bold is True
        assert str(s.font.color.rgb) == "1F3A5F"
        assert s.paragraph_format.page_break_before is False
        # Keep-with-next prevents the subheading from being orphaned at the
        # bottom of a page with its content starting on the next page.
        assert s.paragraph_format.keep_with_next is True
    assert by_name["Heading 5"].font.italic is True

    root = _styles_root(built_docx)
    for display_name in ("heading 4", "heading 5"):
        st = _find_style(root, display_name)
        assert st is not None
        pPr = st.find(f"{{{W}}}pPr")
        assert pPr is None or pPr.find(f"{{{W}}}outlineLvl") is None


def test_req_id_and_assertions_label_styles_defined(built_docx):
    from docx import Document
    d = Document(built_docx)
    by_name = {s.name: s for s in d.styles}
    assert "REQ ID" in by_name
    assert "Assertions Label" in by_name
    assert by_name["REQ ID"].font.name == "Consolas"
    assert by_name["Assertions Label"].font.bold is True
    assert by_name["Assertions Label"].paragraph_format.keep_with_next is True


def test_table_style_has_gray_borders_and_blue_header_row(built_docx):
    root = _styles_root(built_docx)
    table = _find_style(root, "Table")
    assert table is not None

    tblPr = table.find(f"{{{W}}}tblPr")
    borders = tblPr.find(f"{{{W}}}tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = borders.find(f"{{{W}}}{side}")
        assert b is not None, f"missing {side} border"
        assert b.get(f"{{{W}}}color") == "BFBFBF"

    first_row = None
    for tsp in table.findall(f"{{{W}}}tblStylePr"):
        if tsp.get(f"{{{W}}}type") == "firstRow":
            first_row = tsp
            break
    assert first_row is not None
    rPr = first_row.find(f"{{{W}}}rPr")
    assert rPr.find(f"{{{W}}}b") is not None, "first-row should be bold"
    shd = first_row.find(f"{{{W}}}tcPr").find(f"{{{W}}}shd")
    assert shd.get(f"{{{W}}}fill") == "DAEEF3"


def test_footer_layout_table_has_no_visible_borders(built_docx):
    """The footer uses a 3-col table to position left/center/right labels;
    it should NOT carry the global Table style's gray borders."""
    with zipfile.ZipFile(built_docx) as z:
        # There's only one footer in the URS-style docx (footer1.xml).
        footer = etree.fromstring(z.read("word/footer1.xml"))
    tbl = footer.find(f".//{{{W}}}tbl")
    assert tbl is not None, "footer should contain a layout table"
    borders = tbl.find(f"{{{W}}}tblPr/{{{W}}}tblBorders")
    assert borders is not None, "footer table should explicitly suppress borders"
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = borders.find(f"{{{W}}}{side}")
        assert b is not None and b.get(f"{{{W}}}val") == "nil", (
            f"footer table {side} border should be nil, got "
            f"{b.get(f'{{{W}}}val') if b is not None else 'missing'}"
        )


def test_settings_request_field_update_on_open(built_docx):
    with zipfile.ZipFile(built_docx) as z:
        settings = etree.fromstring(z.read("word/settings.xml"))
    update = settings.find(f"{{{W}}}updateFields")
    assert update is not None, "missing <w:updateFields> in settings.xml"
    assert update.get(f"{{{W}}}val") == "true"


def test_cover_styles_centered_with_expected_spacing(built_docx):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    d = Document(built_docx)
    by_name = {s.name: s for s in d.styles}
    sub = by_name["Subtitle"]
    assert sub.font.size.pt == 15
    assert sub.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert sub.paragraph_format.space_before.pt == 12
    abstract = by_name["Abstract"]
    assert abstract.font.size.pt == 10
    assert abstract.paragraph_format.space_before.pt == 15
