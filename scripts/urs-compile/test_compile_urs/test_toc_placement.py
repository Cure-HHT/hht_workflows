"""Integration test for the post-pandoc TOC relocation.

Builds a tiny markdown with cover content + a Heading 1, runs pandoc with
``--toc``, then exercises ``_move_toc_after_cover`` and asserts the SDT now
sits immediately before the first ``Heading 1`` (and the cover paragraphs
are still ahead of it)."""
from __future__ import annotations

import importlib.util
import subprocess
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _load_compile_urs():
    spec_path = Path(__file__).parents[1] / "compile-urs.py"
    spec = importlib.util.spec_from_file_location("compile_urs", spec_path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


@pytest.fixture
def pandoc_docx(tmp_path):
    md = tmp_path / "doc.md"
    md.write_text(
        "**Cover Title Line**\n\n"
        "Cover sub-line one.\n\n"
        "Cover sub-line two.\n\n"
        "# First Chapter\n\n"
        "Body of first chapter.\n\n"
        "## Section\n\n"
        "Section body.\n\n"
        "# Second Chapter\n\n"
        "More body.\n"
    )
    out = tmp_path / "doc.docx"
    try:
        subprocess.run(
            ["pandoc", str(md), "-o", str(out), "--toc",
             "--top-level-division=chapter"],
            check=True, capture_output=True,
        )
    except FileNotFoundError:
        pytest.skip("pandoc not installed")
    return out


def _body_children_summary(docx_path: Path):
    doc = Document(docx_path)
    body = doc.element.body
    out = []
    for child in body:
        tag = child.tag.replace(f"{{{W}}}", "w:")
        if tag == "w:p":
            ps = child.find(f"{qn('w:pPr')}/{qn('w:pStyle')}")
            style_val = ps.get(qn("w:val")) if ps is not None else "(none)"
            text = "".join(t.text or "" for t in child.iter(qn("w:t")))[:40]
            out.append(("p", style_val, text))
        elif tag == "w:sdt":
            out.append(("sdt", "", ""))
        else:
            out.append((tag, "", ""))
    return out


def test_toc_lands_immediately_before_first_heading1(pandoc_docx):
    mod = _load_compile_urs()

    before = _body_children_summary(pandoc_docx)
    # Pandoc baseline: SDT (TOC) is the first non-trivial body child.
    sdt_idx = next(i for i, (t, _, _) in enumerate(before) if t == "sdt")
    h1_idx = next(i for i, (t, s, _) in enumerate(before)
                  if t == "p" and s == "Heading1")
    assert sdt_idx < h1_idx, "baseline expectation: pandoc emits TOC before H1"

    mod._move_toc_after_cover(pandoc_docx)

    after = _body_children_summary(pandoc_docx)
    sdt_idx_after = next(i for i, (t, _, _) in enumerate(after) if t == "sdt")
    h1_idx_after = next(i for i, (t, s, _) in enumerate(after)
                        if t == "p" and s == "Heading1")
    assert sdt_idx_after == h1_idx_after - 1, (
        "expected SDT to sit immediately before first H1; got "
        f"sdt at {sdt_idx_after}, H1 at {h1_idx_after}"
    )

    # Cover paragraphs still precede the TOC.
    cover_paragraphs_before_toc = [
        s for i, (t, s, _) in enumerate(after) if t == "p" and i < sdt_idx_after
    ]
    assert any(s != "Heading1" for s in cover_paragraphs_before_toc), (
        "cover paragraphs should remain ahead of the TOC"
    )


def test_strip_unreferenced_bookmarks_drops_orphans_keeps_referenced(tmp_path):
    """Orphan bookmarks emitted by pandoc on heading anchors should be
    removed; bookmarks targeted by a hyperlink or REF field must survive."""
    mod = _load_compile_urs()
    md = tmp_path / "doc.md"
    md.write_text(
        "# Chapter One {#anchor-orphan}\n\n"
        "Body text.\n\n"
        "# Chapter Two {#anchor-kept}\n\n"
        "See [Chapter Two](#anchor-kept) for details.\n"
    )
    out = tmp_path / "doc.docx"
    try:
        subprocess.run(
            ["pandoc", str(md), "-o", str(out),
             "--top-level-division=chapter"],
            check=True, capture_output=True,
        )
    except FileNotFoundError:
        pytest.skip("pandoc not installed")

    before_names = {b.get(qn("w:name")) for b in
                    Document(out).element.body.iter(qn("w:bookmarkStart"))}
    assert "anchor-orphan" in before_names
    assert "anchor-kept" in before_names

    mod._strip_unreferenced_bookmarks(out)

    after_names = {b.get(qn("w:name")) for b in
                   Document(out).element.body.iter(qn("w:bookmarkStart"))}
    assert "anchor-orphan" not in after_names, "orphan should be stripped"
    assert "anchor-kept" in after_names, (
        "bookmark targeted by [text](#anchor) must survive"
    )


def test_no_op_when_no_toc_present(tmp_path):
    mod = _load_compile_urs()
    md = tmp_path / "doc.md"
    md.write_text("# Title\n\nBody.\n")
    out = tmp_path / "doc.docx"
    try:
        subprocess.run(["pandoc", str(md), "-o", str(out)], check=True,
                       capture_output=True)
    except FileNotFoundError:
        pytest.skip("pandoc not installed")
    before = _body_children_summary(out)
    mod._move_toc_after_cover(out)
    after = _body_children_summary(out)
    assert before == after, "should be a no-op when no TOC SDT is present"
