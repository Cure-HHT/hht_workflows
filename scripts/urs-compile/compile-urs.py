#!/usr/bin/env python3
"""Compile the URS PDF via a single pandoc pass.

Pipeline:
  1. Read elspais graph JSON (build/graph.json by default).
  2. Read tools/urs-section-map.yaml manifest.
  3. Assemble markdown chapter-by-chapter, section-by-section.
     - Emit URS chapter and section headings from manifest.
     - For each section, emit the file prose (REMAINDERs) followed by the
       section's REQs in source order, with PRD/GUI twins sharing one
       kebab name merged into a single level-3 section (see
       urs_compile/ordering.py). Core chapters emit DIARY-* REQs; the
       sponsor chapter collects the sponsor-namespace REQs from the
       files it references.
     - Every body section heading is preceded by a page-break marker
       (raw {=latex} for the PDF, raw {=openxml} for the docx) so each
       level-2 section starts on a fresh page while the first level-3
       heading of the section shares that page.
  4. Prepend frontmatter, append appendices + glossary.
  5. Invoke pandoc once with the URS LaTeX template + cover + resource path.
"""

from __future__ import annotations

import argparse
import re
import subprocess
import sys
from pathlib import Path

import yaml

# Allow `python compile-urs.py` from any cwd (the helper modules live
# alongside this script).
sys.path.insert(0, str(Path(__file__).parent))

import build_docx_reference  # noqa: E402
from urs_compile.graph_loader import Graph  # noqa: E402
from urs_compile.ordering import (  # noqa: E402
    grouped_section_requirements,
    section_remainders,
)
from urs_compile.manifest import Manifest  # noqa: E402
from urs_compile.render import RenderConfig, render_node  # noqa: E402


_HEADING_RE = re.compile(r"^(#{1,5})(\s+)", re.MULTILINE)
_LEADING_H1_RE = re.compile(r"\A\s*#\s+[^\n]*(\n+|\Z)")
_LEADING_COMMENTS_RE = re.compile(r"\A(?:\s*<!--.*?-->\s*\n?)+", re.DOTALL)
_LATEX_BLOCK_RE = re.compile(r"```\{=latex\}\n.*?\n```\n?", re.DOTALL)
# elspais-generated glossary entries pack the "Defined in: <REQ-id>"
# attribution onto the same source line as the definition, separated by
# an inline raw-LaTeX span (`\\*\null\hfill`) that the PDF target uses to
# right-align the attribution. For docx, that span is dropped — leaving
# the attribution mashed onto the definition. The line-break regex
# below substitutes a `<br>` so pandoc emits a `<w:br/>` in the same
# Definition paragraph. The smallskip regex strips the trailing vertical
# spacer (PDF-only).
_INLINE_LATEX_LINEBREAK_RE = re.compile(r"`\\\\\*\\null\\hfill`\{=latex\}")
_INLINE_LATEX_SMALLSKIP_RE = re.compile(r"`\\smallskip`\{=latex\}[ \t]*")
# Force "See:" / "Avoid:" / "Defined in:" labels onto their own lines in
# the docx output. pandoc treats single source newlines as spaces inside
# a definition entry, which collapses these labels into one run-on
# paragraph; injecting an explicit <br> keeps them visually separated.
_DEFLIST_LABEL_RE = re.compile(r"\n(See:|Avoid:|Defined in:)")
_H2_HEADING_RE = re.compile(r"^(## )(?!.*\{\.)(.+?)$", re.MULTILINE)

# Emitted before every body section heading. The {=latex} block starts the
# section on a fresh page and raises the fresh-section flag that the LaTeX
# template's \subsection format consumes to SKIP the page break on the
# section's first level-3 heading. The {=openxml} block is the docx
# equivalent — a bare page-break paragraph that _apply_heading_page_breaks
# later folds into pageBreakBefore on the section heading itself. Each
# pandoc writer ignores the other format's raw block.
SECTION_PAGE_BREAK_MD = (
    "\n```{=latex}\n"
    "\\clearpage\\global\\URSsectionfreshtrue\n"
    "```\n"
    "\n```{=openxml}\n"
    '<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n'
    "```\n"
)
_GLOSSARY_LETTER_HEADING_RE = re.compile(r"^## ([A-Z])\s*$", re.MULTILINE)
_DEFINED_IN_RE = re.compile(r"\n\*Defined in: ([^*]*)\*", re.MULTILINE)


def _break_defined_in(text: str) -> str:
    """Render every `*Defined in: <path>*` line in the elspais glossary /
    references on a new line, right-aligned, with a small gap before the
    next entry but no extra vertical space above.

    elspais emits each entry as a definition list followed immediately
    (no blank line) by `*Defined in: <path>*`. Without intervention,
    pandoc treats the line as a continuation of the last `:` definition
    and renders the pointer inline with the URL. A paragraph break
    would force visible vertical space above, which the user rejected.

    Instead, stay inside the same definition-list item and inject:
    - `\\\\*` — forced linebreak, with `*` to forbid a page break
      between the prior content and the pointer (otherwise the pointer
      can land at the top of the next page, orphaned from its term).
    - `\\null\\hfill` — fills the new line so the italic text
      right-aligns against the right margin.
    - `\\smallskip` after the emph — a small gap before the next term
      so adjacent entries aren't visually flush against each other.
    """
    return _DEFINED_IN_RE.sub(
        r" `\\\\*\\null\\hfill`{=latex} *Defined in: \1* `\\smallskip`{=latex}",
        text,
    )


def _fill_missing_glossary_letters(text: str) -> str:
    """Insert `## <Letter>\\n\\n*No terms.*\\n` for every alphabet letter
    that doesn't already appear as a single-letter H2 heading.

    elspais's `glossary` output only emits letter headings that have terms
    under them. A reader scanning the rendered glossary couldn't tell
    whether a missing letter (K, X, Z, ...) was an authoring oversight or
    an intentional gap. Filling in every absent letter with an explicit
    "No terms." marker disambiguates: an empty letter section is a
    statement that the corpus has no terms beginning with that letter.

    Letters that DO have terms are left untouched. Placeholders are
    inserted just before the next existing letter heading; tail letters
    that come after the last existing one (e.g. Z when last is Y) are
    appended at the end of the glossary section (before any following
    H1 chapter such as References).
    """
    matches = list(_GLOSSARY_LETTER_HEADING_RE.finditer(text))
    if not matches:
        return text  # No letter structure detected; nothing to patch.

    out_parts: list[str] = []
    cursor = 0
    last_letter = "@"  # one before "A"
    for match in matches:
        letter = match.group(1)
        # Letters that should appear in the gap between last_letter and letter.
        missing = [chr(c) for c in range(ord(last_letter) + 1, ord(letter))]
        if missing:
            # Emit text up to this heading, then insert placeholders, then
            # the heading itself in the regular flow.
            out_parts.append(text[cursor:match.start()])
            for m in missing:
                out_parts.append(f"## {m}\n\n*No terms.*\n\n")
            cursor = match.start()
        # Continue normal accumulation; the heading line itself is emitted
        # when we close out the loop or hit the next match.
        last_letter = letter

    # Append everything after the last match boundary.
    out_parts.append(text[cursor:])
    body = "".join(out_parts)

    # Tail letters after the last existing letter (e.g. Z when last is Y).
    if last_letter < "Z":
        tail_letters = [chr(c) for c in range(ord(last_letter) + 1, ord("Z") + 1)]
        tail = "".join(f"\n## {m}\n\n*No terms.*\n" for m in tail_letters)
        # Splice tail before the next H1 (References / Term Index) so the
        # placeholders sit inside the glossary chapter. Search starts AFTER
        # the last existing letter heading so the glossary's own `# Glossary`
        # H1 (which sits at the top of the file) is not matched.
        last_letter_match = matches[-1]
        scan_from = last_letter_match.end()
        next_h1 = re.search(r"\n# [A-Z]", body[scan_from:])
        if next_h1:
            insert_at = scan_from + next_h1.start()
            body = body[:insert_at].rstrip() + "\n" + tail + body[insert_at:]
        else:
            body = body.rstrip() + "\n" + tail
    return body


def _strip_latex_blocks(text: str) -> str:
    """Remove raw LaTeX fenced blocks (`{=latex}`) — for non-LaTeX targets."""
    return _LATEX_BLOCK_RE.sub("", text)


def _normalise_glossary_breaks_for_docx(text: str) -> str:
    """Translate PDF-only glossary line-break tricks into docx equivalents.

    elspais's glossary writer hand-codes inline LaTeX to format each entry
    as definition + right-aligned "Defined in" attribution on a single
    source line. In the docx target the LaTeX spans are dropped and
    single newlines collapse to spaces — leaving a run-on Definition
    paragraph. We substitute pandoc's hard-line-break syntax (two
    trailing spaces before the newline) so each label lands on its own
    line inside the same paragraph. A ``<br>`` HTML tag would be silently
    dropped by pandoc's docx writer, so it can't be used here."""
    text = _INLINE_LATEX_LINEBREAK_RE.sub("  \n", text)
    text = _INLINE_LATEX_SMALLSKIP_RE.sub("", text)
    text = _DEFLIST_LABEL_RE.sub(r"  \n\1", text)
    return text


def _exclude_h2_from_toc(text: str) -> str:
    """Mark every H2 heading as `{.unnumbered .unlisted}` so pandoc skips it
    in the TOC and omits the auto-numbering.

    Used on `_generated/term-index.md` where each defined term is an H2 —
    listing 130+ of those in the TOC would bloat the front matter by
    several pages. The `# Term Index` chapter heading is left untouched
    (still numbered, still in TOC); only the per-term entries are hidden.
    The negative lookahead `(?!.*\\{\\.)` avoids double-tagging if a
    heading already carries an attribute block.
    """
    return _H2_HEADING_RE.sub(r"\1\2 {.unnumbered .unlisted}", text)


def _has_leading_h1(text: str) -> bool:
    """Return True if `text` starts with an H1 heading.

    Tolerates leading HTML comments (e.g. elspais-generated banner comments
    in spec/_generated/*.md) and blank lines before the heading.
    """
    cleaned = _LEADING_COMMENTS_RE.sub("", text)
    return bool(_LEADING_H1_RE.match(cleaned))


def demote_headings(md: str, levels: int = 1) -> str:
    """Demote every ATX heading in `md` by `levels` (capped at H6)."""
    def repl(match: re.Match) -> str:
        hashes = match.group(1)
        new = "#" * min(6, len(hashes) + levels)
        return new + match.group(2)
    return _HEADING_RE.sub(repl, md)


def strip_leading_h1(md: str) -> str:
    """Remove the first top-level (`# ...`) heading from `md`.

    Spec files start with a `# File Title` that the URS manifest already
    surfaces as `## X.Y Section Title`. Drop the redundant heading before
    demoting so the section flows straight from URS heading into content.
    """
    return _LEADING_H1_RE.sub("", md, count=1)


def _resolve(primary: Path, associate: Path | None, relative: str) -> Path | None:
    """Resolve a manifest-relative path against primary then associate root.

    Sponsor-overlay prose (cover, frontmatter, sponsor-info) lives in
    the primary repo; platform-generic prose (chapter intros,
    appendices) typically lives in the associate repo. This helper
    looks in primary first and falls back to associate so the section
    map can be authored once with relative paths regardless of which
    repo carries each piece.
    """
    p = primary / relative
    if p.exists():
        return p
    if associate is not None:
        a = associate / relative
        if a.exists():
            return a
    return None


def assemble_markdown(graph: Graph, manifest: Manifest, primary: Path,
                      associate: Path | None,
                      config: RenderConfig = RenderConfig()) -> str:
    """Build the full assembled markdown document body (no frontmatter / appendices)."""
    parts: list[str] = []
    for chapter in manifest.chapters:
        # LaTeX adds the chapter/section numbers; we just emit titles.
        parts.append(f"\n\n# {chapter.title}\n")
        # Chapter-level intro prose, if the manifest references one. The
        # file is plain markdown and is rendered between the chapter
        # heading and its first section heading. Skip silently if the
        # pointed-at file is missing so an unwired chapter doesn't break
        # the build.
        if chapter.intro_file:
            intro_path = _resolve(primary, associate, chapter.intro_file)
            if intro_path is not None:
                parts.append("\n" + intro_path.read_text().rstrip() + "\n")
        prev_section_num: int | None = None
        for section in chapter.sections:
            # Parse the manifest section number ("5.3" -> 3) and inject a raw
            # \setcounter when the manifest skips a slot (e.g. 5.1 -> 5.3 to
            # preserve a deliberate "User Interface" gap). LaTeX numbers
            # sections sequentially otherwise, which would silently re-number
            # 5.3+ to 5.2+ and confuse reviewers comparing to the manifest.
            try:
                section_idx = int(section.number.split(".")[-1])
            except (ValueError, IndexError):
                section_idx = None
            if (
                section_idx is not None
                and prev_section_num is not None
                and section_idx > prev_section_num + 1
            ):
                parts.append(
                    f"\n```{{=latex}}\n\\setcounter{{section}}{{{section_idx - 1}}}\n```\n"
                )
            prev_section_num = section_idx
            # Page-break marker before the section heading: the raw LaTeX
            # \clearpage starts the section on a fresh page and raises the
            # fresh-section flag the template's \subsection format consumes
            # to skip the page break on the section's FIRST level-3 heading.
            # The raw openxml paragraph is the docx equivalent; pandoc's
            # docx writer passes it through and _apply_heading_page_breaks
            # converts it into pageBreakBefore on the heading. Each writer
            # ignores the other format's raw block.
            parts.append(SECTION_PAGE_BREAK_MD)
            parts.append(f"\n## {section.title}\n")
            # Emit file prose (REMAINDERs) first, then the section's REQ
            # groups (PRD/GUI twins sharing a kebab name render as one
            # level-3 section: the primary REQ carries the heading, the
            # rest follow as heading-less blocks). The sponsor chapter
            # skips REMAINDERs — its prose is the body chapters' job; it
            # only collects the sponsor-namespace REQs.
            # We collect (kind, rendered_text) so we can demote only
            # REMAINDER output — REQ output is already at the final heading
            # level (### title / #### subsections) and must not be demoted.
            rendered_chunks: list[tuple[str, str]] = []
            if chapter.scope == "core":
                for rem in section_remainders(graph, section.files):
                    rendered_chunks.append((rem.kind, render_node(rem, graph, config)))
                    rendered_chunks.append(("SEP", "\n"))
            groups = grouped_section_requirements(
                graph, section.files, scope=chapter.scope,
            )
            for group in groups:
                primary_req = group[0]
                rendered_chunks.append(
                    (primary_req.kind, render_node(primary_req, graph, config))
                )
                rendered_chunks.append(("SEP", "\n"))
                for req in group[1:]:
                    rendered_chunks.append((
                        req.kind,
                        render_node(req, graph, config,
                                    primary_label=primary_req.label),
                    ))
                    rendered_chunks.append(("SEP", "\n"))
            emitted_any = bool(groups) or any(
                kind == "REMAINDER" for kind, _ in rendered_chunks
            )
            if not emitted_any:
                parts.append(
                    f"\n*(No content found for {section.number} — manifest references {section.files})*\n"
                )
                continue
            # Drop the redundant `# File Title` heading the first REMAINDER
            # carries from the spec file; URS section heading already labels it.
            first_rem_idx = next(
                (i for i, (k, _) in enumerate(rendered_chunks) if k == "REMAINDER"),
                None,
            )
            if first_rem_idx is not None:
                kind, text = rendered_chunks[first_rem_idx]
                rendered_chunks[first_rem_idx] = (kind, strip_leading_h1(text))
            # Demote spec-file REMAINDER headings by one level so spec H1 ->
            # H2 subsections under the URS section. REQ chunks are already
            # at the correct level (### / ####) and pass through unchanged.
            final_chunks: list[str] = []
            for kind, text in rendered_chunks:
                if kind == "REMAINDER":
                    text = demote_headings(text, levels=1)
                final_chunks.append(text)
            parts.append("".join(final_chunks))
    return "".join(parts)


def _rewrite_image_paths(text: str) -> str:
    """Rewrite spec-relative image paths so pandoc resolves them from repo root.

    Spec files at `spec/prd-*.md` reference images via `./images/`,
    which is correct relative to the spec file but wrong when pandoc reads the
    assembled markdown from `build/`. Strip the leading `./` so the path is
    relative to the repo root with the `spec/` prefix on `--resource-path`.
    """
    return text.replace("./images/", "spec/images/").replace("(images/", "(spec/images/")


def assemble_full_document(
    graph: Graph,
    manifest: Manifest,
    primary: Path,
    associate: Path | None = None,
    target_format: str = "pdf",
    config: RenderConfig = RenderConfig(),
) -> str:
    """Assemble the full markdown including frontmatter, body, appendices, glossary, term-index.

    `target_format` controls format-specific assembly differences:
    - "pdf": cover is supplied separately via pandoc --variable=cover-tex;
      raw `{=latex}` blocks (e.g. \\setcounter for the deliberate 5.2 skip)
      are preserved.
    - "docx": prepend spec/URS-manifest/cover.md (if present) at the
      top of the assembled doc; strip raw `{=latex}` blocks, which Word
      can't render.

    Manifest-referenced prose (cover.md, frontmatter, appendices, etc.)
    is read from `primary` first and falls back to `associate`, so the
    section map's relative paths can resolve content authored in either
    repo.
    """
    chunks: list[str] = []
    # docx prepends a markdown cover; PDF gets its cover via pandoc variable
    if target_format == "docx":
        cover_md = _resolve(primary, associate, "spec/URS-manifest/cover.md")
        if cover_md is not None:
            chunks.append(cover_md.read_text())
    if manifest.frontmatter:
        path = _resolve(primary, associate, manifest.frontmatter)
        if path is not None:
            chunks.append(path.read_text())
    chunks.append(assemble_markdown(graph, manifest, primary, associate, config))
    # Each back-matter section carries an explicit heading label so
    # "term_index" renders as "Term Index" (not "Term_index" from .title()).
    for key, heading in (
        ("appendices", "Appendices"),
        ("glossary", "Glossary"),
        ("term_index", "Term Index"),
    ):
        path_str = getattr(manifest, key)
        if not path_str:
            continue
        full = _resolve(primary, associate, path_str)
        if full is None:
            continue
        text = full.read_text()
        # The term-index file lists every defined term as an H2 (130+ of
        # them). Without this rewrite, pandoc inflates the TOC with every
        # term — flag each H2 as unnumbered + unlisted.
        if key == "term_index":
            text = _exclude_h2_from_toc(text)
        # Glossary letters that have no terms are absent from elspais's
        # output; fill them in with "No terms." placeholders so a reader
        # can distinguish "intentionally empty" from "authoring oversight".
        # Also break `*Defined in: ...*` onto its own paragraph so it
        # doesn't render inline with the entry's URL / last definition.
        if key == "glossary":
            text = _fill_missing_glossary_letters(text)
            text = _break_defined_in(text)
        # Only inject a chapter heading when the file doesn't already
        # provide one; appendices/glossary/term-index commonly start with
        # `# Appendices` / `# Glossary` / `# Term Index`, possibly after
        # an elspais auto-generated comment banner.
        if not _has_leading_h1(text):
            chunks.append(f"\n\n# {heading}\n\n")
        chunks.append(text)
    text = _rewrite_image_paths("\n\n".join(chunks))
    if target_format != "pdf":
        text = _strip_latex_blocks(text)
        text = _normalise_glossary_breaks_for_docx(text)
    return text


def write_latex_sponsor_header(sponsor_info: dict, output_path: Path) -> None:
    """Write a small LaTeX include-in-header file that overrides sponsor-identity macros.

    urs-template.latex declares \\providecommand{\\sponsorName} (etc.) with
    placeholder defaults; this file \\renewcommand's them with the values
    read from sponsor-info.yaml. Empty / missing keys keep the placeholder.
    """
    lines = []
    for key, macro in (
        ("sponsor_name", "sponsorName"),
        ("protocol_number", "protocolNumber"),
        ("protocol_version", "protocolVersion"),
    ):
        val = sponsor_info.get(key)
        if val:
            lines.append(f"\\renewcommand{{\\{macro}}}{{{val}}}")
    output_path.write_text("\n".join(lines) + "\n")


def run_pandoc_pdf(
    markdown_path: Path,
    output_path: Path,
    template: Path,
    cover: Path,
    resource_paths: list[Path],
    sponsor_header: Path | None = None,
    engine: str = "xelatex",
) -> None:
    filters_dir = Path(__file__).parent / "pandoc-filters"
    table_filter = filters_dir / "table-grid.lua"
    assertion_filter = filters_dir / "assertion-label-italic.lua"
    image_filter = filters_dir / "image-normalize.lua"
    code_filter = filters_dir / "code-breakable.lua"
    cmd = [
        "pandoc",
        str(markdown_path),
        "-o", str(output_path),
        # autolink_bare_uris wraps `https://...` text in `\url{...}` so xurl
        # can break it at any character (without it, pandoc emits the URL
        # as plain text and long URLs overflow the right margin — round-2
        # audit caught this on the JAMA and ASH Publications links in
        # DIARY-PRD-questionnaire-nose-hht / -hht-qol assertion B).
        "-f", "markdown+autolink_bare_uris",
        "--pdf-engine", engine,
        "--template", str(template),
        # Cover is consumed by the template's `$cover-tex$` slot, which
        # places it before the TOC and applies `\thispagestyle{empty}`.
        # `--include-before-body` would land it BETWEEN TOC and body, hiding
        # the cover behind the contents page; use `--variable` instead.
        f"--variable=cover-tex:{cover}",
        "--toc",
        "--toc-depth=3",
        # report class: map `#` -> \chapter so URS chapter numbering (4, 5, 6)
        # survives. pandoc 2.x defaults to \section without this flag, which
        # collapses our chapter headings down a level and yields 0.x numbering.
        "--top-level-division=chapter",
        # Lua filter: re-render every pipe-table as a longtable with full
        # grid lines + shaded header row. LaTeX target only; the filter
        # passes other formats through unchanged.
        f"--lua-filter={table_filter}",
        # Lua filter: italicise A./B./C. labels on alphabetic ordered
        # lists (assertion blocks). LaTeX target only.
        f"--lua-filter={assertion_filter}",
        # Lua filter: normalise every Image to a fixed width and force
        # in-place placement (no float). LaTeX target only.
        f"--lua-filter={image_filter}",
        # Lua filter: inject \allowbreak after every hyphen inside inline
        # `\texttt{}` spans so long REQ ids (DIARY-PRD-...-...) can wrap
        # at hyphens instead of overflowing the right margin.
        f"--lua-filter={code_filter}",
        "--resource-path=" + ":".join(str(p) for p in resource_paths),
    ]
    if sponsor_header is not None:
        cmd.append(f"--include-in-header={sponsor_header}")
    subprocess.run(cmd, check=True)


def _strip_unreferenced_bookmarks(docx_path: Path) -> None:
    """Remove pandoc-emitted bookmarks that no hyperlink or field targets.

    Pandoc inserts a ``<w:bookmarkStart>`` / ``<w:bookmarkEnd>`` pair at
    every heading anchor (slugified title or ``{#req-id}`` block ID). With
    Word's "Show bookmarks" option enabled these appear as gray brackets
    in the body. Word's TOC field generates its own ``_Toc*`` bookmarks
    on update, so pandoc's named ones are only useful when something
    actually targets them (internal hyperlink ``w:anchor=`` or a
    ``REF``/``PAGEREF`` field). The URS body currently has none, so we
    drop the orphans. Anything *referenced* is preserved verbatim so
    future renderer changes that emit cross-references keep working."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    body = doc.element.body

    referenced: set[str] = set()
    for hl in body.iter(qn("w:hyperlink")):
        anchor = hl.get(qn("w:anchor"))
        if anchor:
            referenced.add(anchor)
    for instr in body.iter(qn("w:instrText")):
        text = (instr.text or "").strip()
        # Field-instruction syntax is e.g. ` REF MyBookmark \h ` — split
        # on whitespace and treat every non-switch token as a potential
        # name. Cheap, but conservative: false positives keep a bookmark
        # alive needlessly; false negatives delete a referenced one.
        for tok in text.split():
            if tok and not tok.startswith("\\"):
                referenced.add(tok)

    removed = 0
    for el in list(body.iter(qn("w:bookmarkStart"))):
        name = el.get(qn("w:name"))
        if name in referenced:
            continue
        bid = el.get(qn("w:id"))
        el.getparent().remove(el)
        # Drop the matching bookmarkEnd (paired by id).
        for end in list(body.iter(qn("w:bookmarkEnd"))):
            if end.get(qn("w:id")) == bid:
                end.getparent().remove(end)
                break
        removed += 1

    if removed:
        doc.save(docx_path)


def _apply_heading_page_breaks(docx_path: Path) -> None:
    """Fold assembler page-break markers into heading pagination.

    The assembler emits a raw-openxml paragraph (a bare page-type
    ``<w:br>``) before every body section heading. Leaving it as-is
    would work but strands an empty paragraph at the top of each new
    page, so instead: delete each marker and set ``pageBreakBefore`` on
    the section heading that follows it. pandoc interposes body-level
    ``bookmarkStart``/``bookmarkEnd`` elements between the marker and
    the heading, so the scan skips every non-paragraph sibling.

    Then, for each marked section, clear ``pageBreakBefore`` on the
    FIRST Heading-3 paragraph inside it (per-paragraph ``w:val=0``
    overrides the Heading 3 style's page-break-before) so the section
    heading and its first level-3 heading share a page. Subsequent
    Heading-3 paragraphs keep the style break. Unmarked regions
    (frontmatter, appendices, glossary) are untouched."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    doc = Document(docx_path)
    body = doc.element.body

    def _style(p_el):
        ps = p_el.find(f"{qn('w:pPr')}/{qn('w:pStyle')}")
        return ps.get(qn("w:val")) if ps is not None else None

    def _is_marker(p_el) -> bool:
        if any((t.text or "").strip() for t in p_el.iter(qn("w:t"))):
            return False
        return any(
            br.get(qn("w:type")) == "page" for br in p_el.iter(qn("w:br"))
        )

    def _next_paragraph(el):
        el = el.getnext()
        while el is not None and el.tag != qn("w:p"):
            el = el.getnext()
        return el

    removed = 0
    marked_headings: list = []
    for p_el in list(body.iterchildren(qn("w:p"))):
        if not _is_marker(p_el):
            continue
        heading = _next_paragraph(p_el)
        body.remove(p_el)
        removed += 1
        if heading is not None:
            Paragraph(heading, None).paragraph_format.page_break_before = True
            marked_headings.append(heading)

    # NOTE: lxml element proxies have unstable Python identity, so the
    # first-H3 scan walks forward from each kept heading element rather
    # than re-iterating the body and comparing ids.
    for heading in marked_headings:
        el = heading.getnext()
        while el is not None:
            if el.tag == qn("w:p"):
                style = _style(el)
                if style in ("Heading1", "Heading2"):
                    break
                if style == "Heading3":
                    Paragraph(el, None).paragraph_format.page_break_before = False
                    break
            el = el.getnext()

    if removed:
        doc.save(docx_path)


def _move_toc_after_cover(docx_path: Path) -> None:
    """Relocate pandoc's auto-generated TOC so it lands on its own page,
    after the cover content.

    Pandoc places the ``--toc`` SDT at body index 0. With ``TOCHeading``
    inheriting page-break-before from ``Heading 1`` in our reference doc,
    Word suppresses that page break when the TOC is the very first body
    element — TOC and cover end up sharing page 1. Moving the SDT to
    immediately before the first ``Heading 1`` paragraph fixes the layout:
    cover on page 1, TOC on page 2 (page-break-before now fires), first
    chapter on page 3 (its own page-break-before)."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    body = doc.element.body

    toc_sdt = None
    for child in body:
        if child.tag != qn("w:sdt"):
            continue
        for ps in child.iter(qn("w:pStyle")):
            if ps.get(qn("w:val")) == "TOCHeading":
                toc_sdt = child
                break
        if toc_sdt is not None:
            break
    if toc_sdt is None:
        return

    first_h1 = None
    for child in body:
        if child.tag != qn("w:p"):
            continue
        ps = child.find(f"{qn('w:pPr')}/{qn('w:pStyle')}")
        if ps is not None and ps.get(qn("w:val")) == "Heading1":
            first_h1 = child
            break
    if first_h1 is None:
        return

    body.remove(toc_sdt)
    first_h1.addprevious(toc_sdt)
    doc.save(docx_path)


def run_pandoc_docx(
    markdown_path: Path,
    output_path: Path,
    resource_paths: list[Path],
    reference_doc: Path | None = None,
) -> None:
    """Invoke pandoc to produce a Word .docx output.

    No --template / --pdf-engine — docx uses a reference doc for styling.
    If reference_doc is provided and exists, pandoc uses it for fonts,
    heading styles, header/footer; otherwise pandoc's default styling applies.
    """
    filters_dir = Path(__file__).parent / "pandoc-filters"
    image_filter = filters_dir / "image-normalize.lua"
    table_filter = filters_dir / "table-autofit-docx.lua"
    cmd = [
        "pandoc",
        str(markdown_path),
        "-o", str(output_path),
        "--toc",
        "--toc-depth=3",
        "--number-sections",
        # Match the PDF's chapter-class mapping so headings line up
        # numerically between the two outputs.
        "--top-level-division=chapter",
        # Lua filter: cap every Image's width to a uniform docx-friendly
        # size. The same filter sets fig-pos=H for the LaTeX target;
        # other format branches inside the filter pass through unchanged.
        f"--lua-filter={image_filter}",
        # Lua filter: content-based column widths for auto-width tables
        # (longest-word floor, prose columns absorb the slack) so short
        # identifiers and headers stop wrapping. docx target only;
        # authored-width grid tables pass through unchanged.
        f"--lua-filter={table_filter}",
        "--resource-path=" + ":".join(str(p) for p in resource_paths),
    ]
    if reference_doc is not None and reference_doc.exists():
        cmd.extend(["--reference-doc", str(reference_doc)])
    subprocess.run(cmd, check=True)

    _apply_heading_page_breaks(output_path)
    _move_toc_after_cover(output_path)
    _strip_unreferenced_bookmarks(output_path)


def main() -> int:
    script_dir = Path(__file__).resolve().parent
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--graph", type=Path, default=Path("build/graph.json"))
    p.add_argument(
        "--manifest", type=Path, default=script_dir / "urs-section-map.yaml",
        help="Section ordering manifest (defaults to the one bundled with the script).",
    )
    p.add_argument(
        "--format", choices=["pdf", "docx", "both"], default="both",
        help="Output format(s) to produce (default: both)",
    )
    p.add_argument("--output-md", type=Path, default=Path("build/urs-assembled.md"))
    p.add_argument("--output-pdf", type=Path, default=Path("docs/urs-compiled.pdf"))
    p.add_argument("--output-docx", type=Path, default=Path("docs/urs-compiled.docx"))
    p.add_argument(
        "--template", type=Path, default=script_dir / "urs-template.latex",
        help="LaTeX template (defaults to the one bundled with the script).",
    )
    p.add_argument(
        "--cover", type=Path, default=Path("spec/URS-manifest/urs-cover.tex"),
        help="LaTeX cover snippet read from the primary repo (consumer-supplied).",
    )
    p.add_argument(
        "--sponsor-info", type=Path,
        default=Path("spec/URS-manifest/sponsor-info.yaml"),
        help="YAML file in the primary repo with keys: sponsor_name, "
             "protocol_number, protocol_version. Used to substitute sponsor "
             "identity into the LaTeX template and docx reference.",
    )
    p.add_argument(
        "--associate-root", type=Path, default=None,
        help="Path to the sibling repo's worktree (provides cross-referenced "
             "spec files and images via pandoc's --resource-path). Optional.",
    )
    p.add_argument(
        "--skip-render", action="store_true",
        help="Stop after markdown assembly (useful for tests)",
    )
    # Back-compat: --skip-pdf used to mean the same thing.
    p.add_argument(
        "--skip-pdf", dest="skip_render", action="store_true",
        help=argparse.SUPPRESS,
    )
    args = p.parse_args()

    repo_root = Path(".").resolve()
    args.output_md.parent.mkdir(parents=True, exist_ok=True)
    args.output_pdf.parent.mkdir(parents=True, exist_ok=True)
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)

    # Load sponsor identity (empty dict if no sponsor-info.yaml — placeholders
    # in the template will render visibly so the gap is obvious).
    sponsor_info: dict = {}
    if args.sponsor_info.exists():
        sponsor_info = yaml.safe_load(args.sponsor_info.read_text()) or {}
    render_config = RenderConfig.from_sponsor_info(sponsor_info)

    graph = Graph.from_json_path(args.graph)
    manifest = Manifest.from_yaml_path(args.manifest)

    formats = ["pdf", "docx"] if args.format == "both" else [args.format]

    associate_root = (
        args.associate_root.resolve()
        if args.associate_root is not None and args.associate_root.exists()
        else None
    )

    # Assemble per-format markdown. When producing both formats, write
    # separate `.<fmt>.md` files so each is independently inspectable.
    per_format_md: dict[str, Path] = {}
    for fmt in formats:
        md = assemble_full_document(
            graph, manifest, repo_root, associate_root,
            target_format=fmt, config=render_config,
        )
        if len(formats) > 1:
            md_path = args.output_md.with_suffix(f".{fmt}.md")
        else:
            md_path = args.output_md
        md_path.write_text(md)
        per_format_md[fmt] = md_path
        print(f"Assembled markdown ({fmt}): {md_path} ({len(md):,} chars)", file=sys.stderr)

    if args.skip_render:
        return 0

    resource_paths = [repo_root, repo_root / "spec" / "images"]
    if associate_root:
        resource_paths.append(associate_root)
        associate_images = associate_root / "spec" / "images"
        if associate_images.exists():
            resource_paths.append(associate_images)

    if "pdf" in formats:
        # Build the LaTeX include-in-header that overrides sponsor macros.
        sponsor_header = args.output_md.parent / "sponsor-header.tex"
        write_latex_sponsor_header(sponsor_info, sponsor_header)
        run_pandoc_pdf(
            markdown_path=per_format_md["pdf"],
            output_path=args.output_pdf,
            template=args.template,
            cover=args.cover,
            resource_paths=resource_paths,
            sponsor_header=sponsor_header,
        )
        print(f"Compiled PDF: {args.output_pdf}", file=sys.stderr)
    if "docx" in formats:
        # Regenerate the docx style reference at build time so headers/footers
        # carry the right sponsor identity.
        reference_doc = args.output_md.parent / "urs-reference.docx"
        build_docx_reference.build_reference_docx(reference_doc, sponsor_info)
        run_pandoc_docx(
            markdown_path=per_format_md["docx"],
            output_path=args.output_docx,
            resource_paths=resource_paths,
            reference_doc=reference_doc,
        )
        print(f"Compiled DOCX: {args.output_docx}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
