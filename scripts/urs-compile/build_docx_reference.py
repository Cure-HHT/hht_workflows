#!/usr/bin/env python3
"""Generate a pandoc reference.docx with URS-specific styles.

The reference doc supplies pandoc with the styles to use when emitting
the URS in Word format:

- Title (cover): centered, 24pt bold
- Heading 1 (chapters): bold, 18pt, page-break-before
- Heading 2 (sections): bold, 14pt
- Heading 3 (REQ headings): bold, 12pt, page-break-before (each REQ on
  its own page — matches the LaTeX template's \\newpage on \\subsection)
- Heading 4 (Overview / Assertions / Rationale): bold, URS-blue, 11pt
- Heading 5 (Trigger / Suppression subgroups): bold italic, URS-blue, 10pt
  Both are kept out of Word's outline (outlineLvl cleared); render.py
  also tags each REQ-internal H4 with {.unnumbered} so pandoc skips
  the 6.1.8.x section prefix.
- Body Text: 11pt Arial (Calibri fallback)
- Subtitle / Author / Date / Abstract: cover-page spacing + alignment
- REQ ID, Assertions Label: custom paragraph styles available in the
  styles palette (defined but not auto-applied by the renderer)
- Table: light-gray borders, modest cell padding, bold pale-cyan header row
- Page header: sponsor / protocol / version, right-aligned
- Page footer: "Specific to Protocol [protocol]    CONFIDENTIAL    Page N of M"

Sponsor identity is read from a YAML file (compile-urs.py passes
spec/URS-manifest/sponsor-info.yaml from the consumer repo).

Usage as CLI:
  python3 build-docx-reference.py --sponsor-info path/to/sponsor-info.yaml \\
                                  --output path/to/urs-reference.docx

Usage as a library:
  from build_docx_reference import build_reference_docx
  build_reference_docx(output_path, sponsor_info_dict)
"""
from __future__ import annotations

import argparse
import subprocess
import tempfile
from pathlib import Path

import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


FOOTER_CENTER = "CONFIDENTIAL"


def _header_text(info: dict) -> str:
    return (
        "eCOA User Requirements Specification\n"
        f"Sponsor: {info.get('sponsor_name', '[SPONSOR_NAME]')}\n"
        f"Protocol: {info.get('protocol_number', '[PROTOCOL]')}    "
        f"Version: {info.get('protocol_version', '1.0')}"
    )


def _footer_left(info: dict) -> str:
    return f"Specific to Protocol {info.get('protocol_number', '[PROTOCOL]')}"

# Body font preference: Arial (matches the URS LaTeX template), falling
# back to Calibri (Word default) if Arial isn't available on the renderer.
BODY_FONT = "Arial"


def _add_page_number_field(paragraph) -> None:
    """Insert a 'Page X of Y' field into the given paragraph."""
    run = paragraph.add_run("Page ")
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText_page = OxmlElement("w:instrText")
    instrText_page.set(qn("xml:space"), "preserve")
    instrText_page.text = " PAGE "
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instrText_page)
    run._r.append(fldChar_sep)
    run._r.append(fldChar_end)

    paragraph.add_run(" of ")

    run2 = paragraph.add_run()
    fldChar_begin2 = OxmlElement("w:fldChar")
    fldChar_begin2.set(qn("w:fldCharType"), "begin")
    instrText_pages = OxmlElement("w:instrText")
    instrText_pages.set(qn("xml:space"), "preserve")
    instrText_pages.text = " NUMPAGES "
    fldChar_sep2 = OxmlElement("w:fldChar")
    fldChar_sep2.set(qn("w:fldCharType"), "separate")
    fldChar_end2 = OxmlElement("w:fldChar")
    fldChar_end2.set(qn("w:fldCharType"), "end")
    run2._r.append(fldChar_begin2)
    run2._r.append(instrText_pages)
    run2._r.append(fldChar_sep2)
    run2._r.append(fldChar_end2)


def _get_style(doc, style_name: str):
    """Look up a style by display name (workaround for python-docx 1.2.0
    where `doc.styles["Heading 1"]` raises KeyError; iterate instead)."""
    for s in doc.styles:
        if s.name == style_name:
            return s
    raise KeyError(f"no style with display name {style_name!r}")


def _set_heading_style(doc, style_name: str, size_pt: int, bold: bool = True,
                       italic: bool = False, page_break_before: bool = False,
                       color: tuple[int, int, int] = (0x1F, 0x3A, 0x5F)) -> None:
    """Configure a Heading N style for size, weight, color, and pagination."""
    style = _get_style(doc, style_name)
    font = style.font
    font.name = BODY_FONT
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor(*color)
    pf = style.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    if page_break_before:
        pf.page_break_before = True


def _set_subheading_style(doc, style_name: str, *, size_pt: int = 11,
                          italic: bool = False) -> None:
    """Configure Heading 4 / Heading 5 for use as REQ-internal subheadings
    (Overview / Assertions / Rationale / Trigger / Suppression).

    Bold + URS-blue, smaller than the chapter/section headings, no
    page-break, and with ``outlineLvl`` cleared so the style stops
    participating in Word's document outline / TOC. Pandoc still maps
    ``####`` / ``#####`` to them, but the renderer also tags each
    REQ-internal occurrence with ``{.unnumbered}`` so pandoc skips its
    ``--number-sections`` prefix and TOC entry."""
    style = _get_style(doc, style_name)
    font = style.font
    font.name = BODY_FONT
    font.size = Pt(size_pt)
    font.bold = True
    font.italic = italic
    font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    pf = style.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    pf.page_break_before = False
    pPr = style.element.find(qn("w:pPr"))
    if pPr is not None:
        for ol in pPr.findall(qn("w:outlineLvl")):
            pPr.remove(ol)


def _set_cover_text_style(doc) -> None:
    """Apply URS cover-page spacing/alignment to pandoc's title-block styles."""
    specs = {
        "Body Text": dict(before=9, after=9),
        "Subtitle": dict(size=15, before=12, after=12, center=True, keep_next=True),
        "Author": dict(center=True, keep_next=True),
        "Date": dict(center=True, keep_next=True),
        "Abstract": dict(size=10, before=15, after=15, keep_next=True),
    }
    for name, spec in specs.items():
        try:
            style = _get_style(doc, name)
        except KeyError:
            continue
        style.font.name = BODY_FONT
        if "size" in spec:
            style.font.size = Pt(spec["size"])
        pf = style.paragraph_format
        if "before" in spec:
            pf.space_before = Pt(spec["before"])
        if "after" in spec:
            pf.space_after = Pt(spec["after"])
        if spec.get("center"):
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if spec.get("keep_next"):
            pf.keep_with_next = True


def _add_custom_paragraph_style(
    doc,
    name: str,
    base_name: str,
    *,
    font_name: str | None = None,
    bold: bool = False,
    before_pt: float = 0,
    after_pt: float = 0,
    keep_next: bool = False,
) -> None:
    """Define a custom paragraph style. Idempotent: re-adding the same name
    is a no-op so the function is safe to call against a docx that already
    carries the style."""
    if any(s.name == name for s in doc.styles):
        return
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    try:
        style.base_style = _get_style(doc, base_name)
    except KeyError:
        pass
    if font_name:
        style.font.name = font_name
    if bold:
        style.font.bold = True
    pf = style.paragraph_format
    if before_pt:
        pf.space_before = Pt(before_pt)
    if after_pt:
        pf.space_after = Pt(after_pt)
    if keep_next:
        pf.keep_with_next = True


def _set_table_style(doc) -> None:
    """Configure the default ``Table`` style with light-gray borders, modest
    cell padding, and a bold pale-cyan first-row treatment.

    python-docx's high-level API doesn't reach ``tblStylePr`` /
    ``tblBorders`` / ``shd``, so we manipulate OXML directly."""
    style = _get_style(doc, "Table")
    style_el = style.element

    # Make the style normally visible (pandoc default flags it semiHidden).
    for tag in ("w:semiHidden", "w:unhideWhenUsed"):
        for el in style_el.findall(qn(tag)):
            style_el.remove(el)

    tblPr = style_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        style_el.append(tblPr)

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "BFBFBF")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    existing_mar = tblPr.find(qn("w:tblCellMar"))
    if existing_mar is not None:
        tblPr.remove(existing_mar)
    tblCellMar = OxmlElement("w:tblCellMar")
    for side, width in (("top", "80"), ("left", "120"), ("bottom", "80"), ("right", "120")):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), width)
        m.set(qn("w:type"), "dxa")
        tblCellMar.append(m)
    tblPr.append(tblCellMar)

    for existing_tsp in style_el.findall(qn("w:tblStylePr")):
        if existing_tsp.get(qn("w:type")) == "firstRow":
            style_el.remove(existing_tsp)
    first_row = OxmlElement("w:tblStylePr")
    first_row.set(qn("w:type"), "firstRow")
    rPr = OxmlElement("w:rPr")
    rPr.append(OxmlElement("w:b"))
    first_row.append(rPr)
    first_row.append(OxmlElement("w:tblPr"))
    tcPr = OxmlElement("w:tcPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "DAEEF3")
    tcPr.append(shd)
    first_row.append(tcPr)
    style_el.append(first_row)


def _enable_update_fields_on_open(doc) -> None:
    """Tell Word/LibreOffice to refresh all fields when the document is
    opened. Without this, pandoc's TOC field renders as a blank gray
    placeholder until the reader right-clicks → Update Field."""
    settings = doc.settings.element
    for existing in settings.findall(qn("w:updateFields")):
        settings.remove(existing)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def _set_body_style(doc) -> None:
    style = _get_style(doc, "Normal")
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)


def _set_title_style(doc) -> None:
    # 'Title' is the style pandoc applies to the document's leading-most
    # paragraph when content carries a `# Title` or custom-style="Title".
    # We size and center it for the cover page experience.
    try:
        title = _get_style(doc, "Title")
    except KeyError:
        return
    title.font.name = BODY_FONT
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(24)


def _strip_table_borders(table) -> None:
    """Override the default ``Table`` style's borders on a specific table.

    Used for the footer layout table — the table is purely a positioning
    device for the three footer labels and shouldn't show the gray borders
    the global Table style applies."""
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _install_header(doc, sponsor_info: dict) -> None:
    """Replace each section's default header with the URS sponsor/protocol block."""
    for section in doc.sections:
        header = section.header
        # Clear existing paragraphs and rebuild with our content.
        for p in list(header.paragraphs):
            p._element.getparent().remove(p._element)
        for line_idx, line in enumerate(_header_text(sponsor_info).split("\n")):
            p = header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(line)
            run.font.name = BODY_FONT
            run.font.size = Pt(9)
            if line_idx == 0:
                run.bold = True


def _install_footer(doc, sponsor_info: dict) -> None:
    """Footer: 'Specific to Protocol ...    CONFIDENTIAL    Page X of Y'.

    python-docx doesn't expose tab stops on default footer paragraphs in
    a clean way; we use a 3-column table so the three labels land left,
    center, and right reliably.
    """
    for section in doc.sections:
        footer = section.footer
        # Clear any existing paragraphs
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)
        table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
        # Suppress borders: the default Table style (intentionally) draws
        # light-gray borders around every cell, but this table is purely
        # for laying out the three footer labels left/center/right.
        _strip_table_borders(table)
        # Left cell
        left = table.cell(0, 0).paragraphs[0]
        left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = left.add_run(_footer_left(sponsor_info))
        run.font.size = Pt(9)
        # Center cell
        center = table.cell(0, 1).paragraphs[0]
        center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = center.add_run(FOOTER_CENTER)
        run.font.size = Pt(9)
        run.bold = True
        # Right cell — page number field
        right = table.cell(0, 2).paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Set the font on the runs we add
        _add_page_number_field(right)
        for r in right.runs:
            r.font.size = Pt(9)


def build_reference_docx(output_path: Path, sponsor_info: dict) -> None:
    """Build a pandoc reference.docx with URS-specific styles and sponsor identity."""
    # Start from pandoc's default reference doc — it ships with all the
    # style names pandoc expects to map markdown elements to.
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        default_path = Path(tmp.name)
    subprocess.run(
        ["pandoc", "--print-default-data-file", "reference.docx"],
        stdout=default_path.open("wb"),
        check=True,
    )

    doc = Document(default_path)

    _set_body_style(doc)
    _set_title_style(doc)
    _set_cover_text_style(doc)
    _set_heading_style(doc, "Heading 1", size_pt=18, page_break_before=True)
    _set_heading_style(doc, "Heading 2", size_pt=14)
    _set_heading_style(doc, "Heading 3", size_pt=12, page_break_before=True)
    _set_subheading_style(doc, "Heading 4", size_pt=11)
    _set_subheading_style(doc, "Heading 5", size_pt=10, italic=True)

    _add_custom_paragraph_style(
        doc, "REQ ID", base_name="Normal",
        font_name="Consolas", before_pt=6, after_pt=4,
    )
    _add_custom_paragraph_style(
        doc, "Assertions Label", base_name="Normal",
        bold=True, before_pt=8, after_pt=4, keep_next=True,
    )

    _set_table_style(doc)
    _enable_update_fields_on_open(doc)

    _install_header(doc, sponsor_info)
    _install_footer(doc, sponsor_info)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate URS-styled reference.docx for pandoc.")
    parser.add_argument(
        "--sponsor-info", type=Path, required=True,
        help="Path to sponsor-info.yaml (keys: sponsor_name, protocol_number, protocol_version).",
    )
    parser.add_argument(
        "--output", type=Path, required=True,
        help="Path to write the generated urs-reference.docx.",
    )
    args = parser.parse_args()

    sponsor_info = yaml.safe_load(args.sponsor_info.read_text())
    build_reference_docx(args.output, sponsor_info)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
